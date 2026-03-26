"""FastAPI backend for Pramana."""

import asyncio
import json
import logging
import threading
import traceback
import uuid
from contextlib import asynccontextmanager
from datetime import datetime, timezone

from fastapi import BackgroundTasks, FastAPI, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from starlette.responses import StreamingResponse

from pramana.config import get_settings
from pramana.llm.sanitize import sanitize_user_input
from pramana.models.database import create_tables, get_engine, get_session, seed_venues
from pramana.models.schema import AnalysisRun, Annotation, ExpertFeedback, Hypothesis, Paper
from pramana.models.schema import ExtractedFact as ExtractedFactDB
from pramana.models.vectors import get_chroma_client, get_evidence_collection, search_evidence

logger = logging.getLogger(__name__)

# In-memory store for running analyses
_analysis_store: dict[str, dict] = {}
_pdf_store: dict[str, dict] = {}  # file_id -> {"text": str, "filename": str, "page_count": int}


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Initialize database on startup."""
    # Configure logging for all pramana modules
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s %(levelname)-8s [%(name)s] %(message)s",
        datefmt="%H:%M:%S",
    )
    # Set pramana loggers to DEBUG for detailed output
    logging.getLogger("pramana").setLevel(logging.DEBUG)
    # Quieten noisy libraries
    logging.getLogger("httpx").setLevel(logging.WARNING)
    logging.getLogger("chromadb").setLevel(logging.WARNING)
    logging.getLogger("openai").setLevel(logging.WARNING)

    settings = get_settings()
    settings.ensure_dirs()
    engine = get_engine(settings)
    create_tables(engine)
    seed_venues(settings)

    # Check if existing DB needs migration for new columns
    try:
        import sqlalchemy
        inspector = sqlalchemy.inspect(engine)
        columns = [c["name"] for c in inspector.get_columns("extracted_facts")]
        if "confidence" not in columns:
            logger.warning(
                "DB migration needed: 'extracted_facts' table is missing 'confidence' column. "
                "Run: ALTER TABLE extracted_facts ADD COLUMN confidence REAL DEFAULT 0.0"
            )
    except Exception:
        pass  # Table may not exist yet — create_tables will handle it

    logger.info("Pramana API started (LLM: %s @ %s)", settings.llm_model, settings.llm_base_url)
    yield


app = FastAPI(
    title="Pramana",
    description="Hypothesis-driven research assistant API",
    version="0.1.0",
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# --- Request/Response models ---

class AnalyzeRequest(BaseModel):
    hypothesis: str
    initiation_type: str = "new"
    max_papers: int = 50
    prior_research: str = ""
    pdf_file_ids: list[str] = []
    domain: str = ""   # User-declared domain (e.g. "Computer Science", "Economics")
    action: str = ""   # Free-text: what the user wants to do (routes to analysis flows)


class AnalyzeResponse(BaseModel):
    run_id: str
    status: str


class ChatMessage(BaseModel):
    role: str  # "user" or "assistant"
    content: str


class ChatRequest(BaseModel):
    message: str
    history: list[ChatMessage] = []


class RunStatus(BaseModel):
    run_id: str
    status: str
    stage: str = ""
    progress: dict = {}
    error: str | None = None


class PaperResponse(BaseModel):
    id: int
    title: str
    authors: list[str]
    year: int | None
    venue: str
    doi: str | None
    abstract: str
    facts: list[dict] = []


class EvidenceSearchResponse(BaseModel):
    results: list[dict]
    total: int


class FeedbackRequest(BaseModel):
    fact_id: int
    action: str  # "confirm", "reject", "comment"
    comment: str = ""


class FeedbackResponse(BaseModel):
    id: int
    fact_id: int
    action: str
    comment: str
    created_at: str


class VenueResponse(BaseModel):
    id: int
    name: str
    venue_type: str
    domain: str
    tier: str


# --- Endpoints ---

@app.post("/api/upload-pdf")
async def upload_pdf(file: UploadFile):
    """Upload a PDF and extract its text for use as prior research context."""
    if not file.filename or not file.filename.lower().endswith(".pdf"):
        raise HTTPException(400, "File must be a PDF")

    content = await file.read()
    if not content:
        raise HTTPException(400, "File is empty")

    if not content[:5] == b"%PDF-":
        raise HTTPException(400, "File does not appear to be a valid PDF")

    from pramana.sources.pdf import extract_text_from_bytes

    text = extract_text_from_bytes(content)
    if not text.strip():
        raise HTTPException(400, "Could not extract text from PDF")

    page_count = text.count("[Page ")

    file_id = str(uuid.uuid4())
    _pdf_store[file_id] = {
        "text": text,
        "filename": file.filename,
        "page_count": page_count,
    }

    return {
        "file_id": file_id,
        "filename": file.filename,
        "page_count": page_count,
        "char_count": len(text),
        "text_preview": text[:500],
    }


@app.post("/api/analyze", response_model=AnalyzeResponse)
async def start_analysis(request: AnalyzeRequest, background_tasks: BackgroundTasks):
    """Start a new analysis run."""
    request.hypothesis = sanitize_user_input(request.hypothesis, max_length=5000)
    request.prior_research = sanitize_user_input(request.prior_research, max_length=5000)

    # Append uploaded PDF text to prior research
    pdf_texts = []
    for fid in request.pdf_file_ids:
        if fid in _pdf_store:
            pdf_data = _pdf_store[fid]
            pdf_texts.append(
                f"--- Uploaded: {pdf_data['filename']} ---\n{pdf_data['text']}"
            )
    if pdf_texts:
        separator = "\n\n"
        combined = separator.join(pdf_texts)
        if len(combined) > 20000:
            combined = combined[:20000]
        if request.prior_research:
            request.prior_research += "\n\n" + combined
        else:
            request.prior_research = combined

    run_id = str(uuid.uuid4())
    settings = get_settings()

    # Create DB records
    with get_session(settings) as session:
        hypothesis = Hypothesis(
            text=request.hypothesis,
            initiation_type=request.initiation_type,
        )
        session.add(hypothesis)
        session.flush()

        run = AnalysisRun(
            hypothesis_id=hypothesis.id,
            status="pending",
            config=json.dumps({
                "max_papers": request.max_papers,
                "initiation_type": request.initiation_type,
            }),
        )
        session.add(run)
        session.flush()
        db_run_id = run.id

    _analysis_store[run_id] = {
        "status": "pending",
        "stage": "queued",
        "db_run_id": db_run_id,
        "hypothesis": request.hypothesis,
        "initiation_type": request.initiation_type,
        "max_papers": request.max_papers,
        "prior_research": request.prior_research,
        "domain": request.domain,
        "action": request.action,
        "progress": {},
        "result": None,
        "error": None,
        # Human-in-the-loop pause points
        "_confirm_event": threading.Event(),   # unblocked by confirm-hypothesis endpoint
        "_curate_event": threading.Event(),    # unblocked by confirm-corpus endpoint
        "parsed_query": None,                  # set after stage 1 pause
        "corpus_papers": None,                 # set after stage 3 pause
    }

    background_tasks.add_task(_run_analysis, run_id)

    return AnalyzeResponse(run_id=run_id, status="pending")


@app.get("/api/analyze/{run_id}", response_model=RunStatus)
async def get_analysis_status(run_id: str):
    """Get analysis run status."""
    if run_id not in _analysis_store:
        raise HTTPException(404, "Analysis run not found")

    store = _analysis_store[run_id]
    return RunStatus(
        run_id=run_id,
        status=store["status"],
        stage=store["stage"],
        progress=store["progress"],
        error=store["error"],
    )


@app.get("/api/analyze/{run_id}/stream")
async def stream_analysis_progress(run_id: str):
    """Stream analysis progress via Server-Sent Events."""
    if run_id not in _analysis_store:
        raise HTTPException(404, "Analysis run not found")

    async def event_generator():
        while True:
            if run_id not in _analysis_store:
                break
            store = _analysis_store[run_id]
            payload = json.dumps({
                "run_id": run_id,
                "status": store["status"],
                "stage": store["stage"],
                "progress": store["progress"],
                "error": store["error"],
            })
            yield f"data: {payload}\n\n"

            if store["status"] in ("completed", "failed",
                                    "awaiting_confirmation", "awaiting_curation"):
                break
            await asyncio.sleep(0.5)

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@app.get("/api/analyze/{run_id}/report")
async def get_analysis_report(run_id: str):
    """Get completed analysis report."""
    # Try in-memory first
    if run_id in _analysis_store:
        store = _analysis_store[run_id]
        if store["status"] != "completed":
            raise HTTPException(400, f"Analysis is {store['status']}, not completed")
        return {"run_id": run_id, "report": store["result"]}

    # Fall back to DB
    settings = get_settings()
    with get_session(settings) as session:
        db_run = session.query(AnalysisRun).filter(
            AnalysisRun.id == _try_parse_int(run_id)
        ).first() if _try_parse_int(run_id) is not None else None
        # Also try matching by UUID stored in memory (check all runs)
        if not db_run or not db_run.results:
            raise HTTPException(404, "Report not found")
        return {"run_id": run_id, "report": json.loads(db_run.results)}


@app.get("/api/papers/{paper_id}", response_model=PaperResponse)
async def get_paper(paper_id: int):
    """Get paper details with extracted evidence."""
    settings = get_settings()
    with get_session(settings) as session:
        paper = session.get(Paper, paper_id)
        if not paper:
            raise HTTPException(404, "Paper not found")

        facts = session.query(ExtractedFactDB).filter_by(paper_id=paper_id).all()
        authors = json.loads(paper.authors) if paper.authors else []

        return PaperResponse(
            id=paper.id,
            title=paper.title,
            authors=authors,
            year=paper.year,
            venue=paper.venue or "",
            doi=paper.doi,
            abstract=paper.abstract or "",
            facts=[
                {
                    "id": f.id,
                    "fact_type": f.fact_type,
                    "content": f.content,
                    "direct_quote": f.direct_quote,
                    "location": f.location,
                    "confidence": f.confidence or 0.0,
                }
                for f in facts
            ],
        )


@app.get("/api/evidence", response_model=EvidenceSearchResponse)
async def search_evidence_endpoint(query: str, limit: int = 20):
    """Search evidence using semantic search."""
    settings = get_settings()
    try:
        chroma = get_chroma_client(settings)
        collection = get_evidence_collection(chroma)
        results = search_evidence(collection, query, n_results=limit)

        items = []
        if results.get("documents"):
            for i, doc in enumerate(results["documents"][0]):
                metadata = results["metadatas"][0][i] if results.get("metadatas") else {}
                distance = results["distances"][0][i] if results.get("distances") else None
                items.append({
                    "text": doc,
                    "metadata": metadata,
                    "score": 1.0 - (distance or 0),
                })

        return EvidenceSearchResponse(results=items, total=len(items))
    except Exception:
        return EvidenceSearchResponse(results=[], total=0)


@app.get("/api/venues", response_model=list[VenueResponse])
async def list_venues(domain: str | None = None):
    """List known research venues."""
    settings = get_settings()
    from pramana.models.schema import Venue

    with get_session(settings) as session:
        query = session.query(Venue)
        if domain:
            query = query.filter(Venue.domain.contains(domain))
        venues = query.all()

        return [
            VenueResponse(
                id=v.id,
                name=v.name,
                venue_type=v.venue_type or "",
                domain=v.domain or "",
                tier=v.tier or "",
            )
            for v in venues
        ]


def _try_parse_int(value: str) -> int | None:
    """Try to parse a string as int, return None if not possible."""
    try:
        return int(value)
    except (ValueError, TypeError):
        return None


@app.get("/api/reports")
async def list_reports():
    """List all completed analysis runs."""
    settings = get_settings()
    with get_session(settings) as session:
        runs = (
            session.query(AnalysisRun)
            .filter(AnalysisRun.status == "completed", AnalysisRun.results.isnot(None))
            .order_by(AnalysisRun.completed_at.desc())
            .all()
        )
        results = []
        for run in runs:
            hypothesis = session.get(Hypothesis, run.hypothesis_id)
            # Extract paper count from stored results
            paper_count = 0
            if run.results:
                try:
                    report = json.loads(run.results)
                    for lr in report.get("lens_results", []):
                        if lr.get("lens") == "evidence_table":
                            paper_count = lr.get("content", {}).get("papers_with_evidence", 0)
                            break
                except (json.JSONDecodeError, KeyError):
                    pass
            results.append({
                "run_id": run.id,
                "hypothesis": hypothesis.text if hypothesis else "",
                "completed_at": run.completed_at.isoformat() if run.completed_at else None,
                "paper_count": paper_count,
            })
        return {"reports": results}


@app.get("/api/reports/{run_id}")
async def get_saved_report(run_id: int):
    """Get a saved report by DB run ID."""
    settings = get_settings()
    with get_session(settings) as session:
        run = session.get(AnalysisRun, run_id)
        if not run or not run.results:
            raise HTTPException(404, "Report not found")
        return {"run_id": run_id, "report": json.loads(run.results)}


@app.post("/api/reports/{run_id}/chat")
async def chat_with_report(run_id: str, request: ChatRequest):
    """Chat about a report using RAG-augmented context."""
    from pramana.llm.client import chat as llm_chat
    from pramana.llm.prompts import REPORT_CHAT_SYSTEM
    from pramana.pipeline.rag import format_retrieved_context, retrieve_relevant_evidence

    settings = get_settings()
    message = sanitize_user_input(request.message, max_length=2000)
    if not message:
        raise HTTPException(400, "Message cannot be empty")

    # Load report — try in-memory first, then DB
    report_data = None
    if run_id in _analysis_store:
        store = _analysis_store[run_id]
        if store["status"] == "completed" and store["result"]:
            report_data = store["result"]

    if report_data is None:
        db_id = _try_parse_int(run_id)
        if db_id is not None:
            with get_session(settings) as session:
                db_run = session.get(AnalysisRun, db_id)
                if db_run and db_run.results:
                    report_data = json.loads(db_run.results)

    if report_data is None:
        raise HTTPException(404, "Report not found")

    # Build report summary (truncated markdown-style)
    report_summary = _build_report_summary(report_data, max_chars=8000)

    # Retrieve relevant evidence via RAG
    rag_results = retrieve_relevant_evidence(message, settings, n_results=20)
    rag_context = format_retrieved_context(rag_results, max_chars=8000)

    # Build conversation messages
    messages: list[dict] = [{"role": "system", "content": REPORT_CHAT_SYSTEM}]

    # Add report context as a system-level context block
    context_block = f"## Report Summary\n{report_summary}\n\n## Retrieved Evidence\n{rag_context}"
    messages.append({"role": "system", "content": context_block})

    # Add conversation history (last 10 messages)
    for msg in request.history[-10:]:
        messages.append({"role": msg.role, "content": msg.content})

    # Add current user message
    messages.append({"role": "user", "content": message})

    try:
        response = llm_chat(messages, settings, temperature=0.3, max_tokens=1500)
        return {"response": response}
    except Exception as e:
        logger.error("Chat LLM call failed: %s", e)
        raise HTTPException(500, "Failed to generate response")


@app.post("/api/feedback", response_model=FeedbackResponse)
async def submit_feedback(request: FeedbackRequest):
    """Submit expert feedback on an extracted fact."""
    if request.action not in ("confirm", "reject", "comment"):
        raise HTTPException(400, "Action must be 'confirm', 'reject', or 'comment'")

    settings = get_settings()
    with get_session(settings) as session:
        fact = session.get(ExtractedFactDB, request.fact_id)
        if not fact:
            raise HTTPException(404, "Fact not found")

        fb = ExpertFeedback(
            fact_id=request.fact_id,
            action=request.action,
            comment=request.comment,
        )
        session.add(fb)
        session.flush()

        return FeedbackResponse(
            id=fb.id,
            fact_id=fb.fact_id,
            action=fb.action,
            comment=fb.comment or "",
            created_at=fb.created_at.isoformat() if fb.created_at else "",
        )


@app.get("/api/feedback/{fact_id}", response_model=list[FeedbackResponse])
async def get_feedback(fact_id: int):
    """Get all feedback for a specific fact."""
    settings = get_settings()
    with get_session(settings) as session:
        feedbacks = (
            session.query(ExpertFeedback)
            .filter(ExpertFeedback.fact_id == fact_id)
            .order_by(ExpertFeedback.created_at.desc())
            .all()
        )
        return [
            FeedbackResponse(
                id=fb.id,
                fact_id=fb.fact_id,
                action=fb.action,
                comment=fb.comment or "",
                created_at=fb.created_at.isoformat() if fb.created_at else "",
            )
            for fb in feedbacks
        ]


@app.get("/api/reports/{run_id}/export")
async def export_report(run_id: int, format: str = "bibtex"):
    """Export report data in BibTeX, CSV, or Markdown format."""
    import csv
    import io

    settings = get_settings()
    with get_session(settings) as session:
        run = session.get(AnalysisRun, run_id)
        if not run:
            raise HTTPException(404, "Report not found")

        # Get all papers associated with this run via extracted facts
        papers = (
            session.query(Paper)
            .join(ExtractedFactDB, ExtractedFactDB.paper_id == Paper.id)
            .filter(
                ExtractedFactDB.paper_id.in_(
                    session.query(ExtractedFactDB.paper_id)
                    .filter(ExtractedFactDB.paper_id == Paper.id)
                )
            )
            .distinct()
            .all()
        )

        if format == "bibtex":
            lines = []
            for p in papers:
                authors_list = json.loads(p.authors) if p.authors else []
                authors_str = " and ".join(authors_list) if authors_list else "Unknown"
                key = _bibtex_key(authors_list, p.year, p.title)
                entry_type = "article"
                lines.append(f"@{entry_type}{{{key},")
                lines.append(f"  title = {{{p.title}}},")
                lines.append(f"  author = {{{authors_str}}},")
                if p.year:
                    lines.append(f"  year = {{{p.year}}},")
                if p.venue:
                    lines.append(f"  journal = {{{p.venue}}},")
                if p.doi:
                    lines.append(f"  doi = {{{p.doi}}},")
                if p.url:
                    lines.append(f"  url = {{{p.url}}},")
                lines.append("}")
                lines.append("")
            content = "\n".join(lines)
            from starlette.responses import Response
            return Response(
                content=content,
                media_type="text/plain",
                headers={"Content-Disposition": f"attachment; filename=pramana_{run_id}.bib"},
            )

        elif format == "csv":
            facts = (
                session.query(ExtractedFactDB, Paper.title)
                .join(Paper, ExtractedFactDB.paper_id == Paper.id)
                .all()
            )
            buf = io.StringIO()
            writer = csv.writer(buf)
            writer.writerow(["paper_title", "fact_type", "content", "direct_quote",
                              "location", "confidence"])
            for fact, paper_title in facts:
                writer.writerow([
                    paper_title,
                    fact.fact_type,
                    fact.content,
                    fact.direct_quote,
                    fact.location,
                    round(fact.confidence or 0.0, 3),
                ])
            from starlette.responses import Response
            return Response(
                content=buf.getvalue(),
                media_type="text/csv",
                headers={"Content-Disposition": f"attachment; filename=pramana_{run_id}_facts.csv"},
            )

        elif format == "markdown":
            if not run.results:
                raise HTTPException(400, "Report not yet completed")
            from pramana.lenses.base import LensResult
            from pramana.pipeline.hypothesis import HypothesisQuery
            from pramana.pipeline.orchestrator import AnalysisResults
            from pramana.report.generator import generate_report

            report_data = json.loads(run.results)
            results = AnalysisResults()
            for lr in report_data.get("lens_results", []):
                results.add(LensResult(
                    lens_name=lr["lens"],
                    title=lr.get("title", lr["lens"]),
                    content=lr.get("content", {}),
                    summary=lr.get("summary", ""),
                ))
            hyp_data = report_data.get("hypothesis", {})
            query = HypothesisQuery(**hyp_data) if hyp_data else HypothesisQuery()
            md = generate_report(results, query, "markdown", settings)
            from starlette.responses import Response
            return Response(
                content=md,
                media_type="text/markdown",
                headers={
                    "Content-Disposition": f"attachment; filename=pramana_{run_id}_report.md"
                },
            )

        elif format == "docx":
            if not run.results:
                raise HTTPException(400, "Report not yet completed")
            from docx import Document
            from starlette.responses import Response as StarResponse

            report_data = json.loads(run.results)
            doc = Document()
            doc.add_heading("Pramana Research Analysis Report", 0)
            hyp = report_data.get("hypothesis", {})
            if hyp.get("topics"):
                doc.add_paragraph(f"Topics: {', '.join(hyp['topics'])}")
            if hyp.get("domains"):
                doc.add_paragraph(f"Domains: {', '.join(hyp['domains'])}")
            doc.add_paragraph("")
            exec_sum = report_data.get("executive_summary", {})
            if isinstance(exec_sum, dict) and exec_sum.get("headline"):
                doc.add_heading("Executive Summary", 1)
                doc.add_paragraph(exec_sum["headline"]).runs[0].bold = True
                for bullet in exec_sum.get("bullets", []):
                    doc.add_paragraph(bullet, style="List Bullet")
            for lr in report_data.get("lens_results", []):
                doc.add_heading(lr.get("title", lr.get("lens", "")), 2)
                doc.add_paragraph(lr.get("summary", "")).runs[0].italic = True
                content = lr.get("content", {})
                if lr.get("lens") == "gap_discovery":
                    for gap in content.get("gaps", []):
                        p = doc.add_paragraph(style="List Bullet")
                        run = p.add_run(f"[{gap.get('severity','?')}] ")
                        run.bold = True
                        p.add_run(gap.get("description", ""))
                elif lr.get("lens") == "lit_review":
                    doc.add_paragraph(content.get("draft", ""))
                elif lr.get("lens") == "research_proposal":
                    for section in ["background", "gap_statement", "methodology"]:
                        val = content.get(section, "")
                        if val:
                            doc.add_heading(section.replace("_", " ").title(), 3)
                            doc.add_paragraph(val)
            import io
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return StarResponse(
                content=buf.read(),
                media_type=(
                    "application/vnd.openxmlformats-officedocument"
                    ".wordprocessingml.document"
                ),
                headers={
                    "Content-Disposition": (
                        f"attachment; filename=pramana_{run_id}_report.docx"
                    )
                },
            )

        else:
            raise HTTPException(400, "format must be 'bibtex', 'csv', 'markdown', or 'docx'")


# --- Human-in-the-loop pause endpoints ---

@app.get("/api/analyze/{run_id}/parsed-query")
async def get_parsed_query(run_id: str):
    """Get the parsed hypothesis query after stage 1 (hypothesis confirmation)."""
    if run_id not in _analysis_store:
        raise HTTPException(404, "Analysis run not found")
    store = _analysis_store[run_id]
    if store.get("parsed_query") is None:
        raise HTTPException(400, "Parsed query not yet available")
    return {"run_id": run_id, "status": store["status"], "parsed_query": store["parsed_query"]}


class ConfirmHypothesisRequest(BaseModel):
    domains: list[str] | None = None
    topics: list[str] | None = None
    search_queries: list[str] | None = None


@app.post("/api/analyze/{run_id}/confirm-hypothesis")
async def confirm_hypothesis(run_id: str, request: ConfirmHypothesisRequest):
    """Confirm or edit the parsed hypothesis to unblock the pipeline."""
    if run_id not in _analysis_store:
        raise HTTPException(404, "Analysis run not found")
    store = _analysis_store[run_id]
    if store["status"] != "awaiting_confirmation":
        raise HTTPException(400, f"Run is not awaiting confirmation (status={store['status']})")
    # Apply edits if provided
    pq = store["parsed_query"]
    if request.domains is not None:
        pq["domains"] = request.domains
    if request.topics is not None:
        pq["topics"] = request.topics
    if request.search_queries is not None:
        pq["search_queries"] = request.search_queries
    store["parsed_query"] = pq
    store["_confirm_edits"] = pq  # pipeline reads this
    store["_confirm_event"].set()
    return {"ok": True}


@app.get("/api/analyze/{run_id}/corpus-papers")
async def get_corpus_papers(run_id: str):
    """Get the paper list after screening (paper curation step)."""
    if run_id not in _analysis_store:
        raise HTTPException(404, "Analysis run not found")
    store = _analysis_store[run_id]
    if store.get("corpus_papers") is None:
        raise HTTPException(400, "Corpus papers not yet available")
    return {
        "run_id": run_id,
        "status": store["status"],
        "papers": store["corpus_papers"],
    }


class ConfirmCorpusRequest(BaseModel):
    excluded_ids: list[int] = []  # db_id values of papers to exclude


@app.post("/api/analyze/{run_id}/confirm-corpus")
async def confirm_corpus(run_id: str, request: ConfirmCorpusRequest):
    """Confirm paper selection and unblock the pipeline for extraction."""
    if run_id not in _analysis_store:
        raise HTTPException(404, "Analysis run not found")
    store = _analysis_store[run_id]
    if store["status"] != "awaiting_curation":
        raise HTTPException(400, f"Run is not awaiting curation (status={store['status']})")
    store["_excluded_paper_ids"] = set(request.excluded_ids)
    store["_curate_event"].set()
    return {"ok": True, "excluded": len(request.excluded_ids)}


# --- Batch I: Annotations, re-run lens, follow-up search ---

class AnnotationRequest(BaseModel):
    content_ref: str   # e.g. "gap:0", "finding:3", "lens:lit_review"
    note: str = ""


@app.post("/api/reports/{run_id}/annotations")
async def create_annotation(run_id: str, request: AnnotationRequest):
    """Bookmark/annotate a finding, gap, or lens result."""
    settings = get_settings()
    with get_session(settings) as session:
        ann = Annotation(run_id=run_id, content_ref=request.content_ref, note=request.note)
        session.add(ann)
        session.flush()
        created_at = ann.created_at.isoformat() if ann.created_at else ""
        return {"id": ann.id, "run_id": run_id, "content_ref": ann.content_ref,
                "note": ann.note, "created_at": created_at}


@app.get("/api/reports/{run_id}/annotations")
async def list_annotations(run_id: str):
    """List all annotations for a run."""
    settings = get_settings()
    with get_session(settings) as session:
        anns = session.query(Annotation).filter(Annotation.run_id == run_id).all()
        return {"annotations": [
            {"id": a.id, "content_ref": a.content_ref, "note": a.note,
             "created_at": a.created_at.isoformat() if a.created_at else ""}
            for a in anns
        ]}


@app.delete("/api/reports/{run_id}/annotations/{ann_id}")
async def delete_annotation(run_id: str, ann_id: int):
    """Delete an annotation."""
    settings = get_settings()
    with get_session(settings) as session:
        ann = session.get(Annotation, ann_id)
        if not ann or ann.run_id != run_id:
            raise HTTPException(404, "Annotation not found")
        session.delete(ann)
        return {"ok": True}


class RerunLensRequest(BaseModel):
    lens_name: str


@app.post("/api/analyze/{run_id}/rerun-lens")
async def rerun_lens(run_id: str, request: RerunLensRequest):
    """Re-run a single lens on the existing corpus/evidence without re-extraction."""
    if run_id not in _analysis_store:
        raise HTTPException(404, "Analysis run not found")
    store = _analysis_store[run_id]
    if store["status"] != "completed":
        raise HTTPException(400, "Run must be completed to re-run a lens")

    settings = get_settings()
    from pramana.pipeline.hypothesis import HypothesisQuery

    # Reconstruct query from stored report
    report_data = store["result"]
    hyp_data = report_data.get("hypothesis", {})
    query = HypothesisQuery(**{k: v for k, v in hyp_data.items()
                               if k in HypothesisQuery.model_fields})

    # Re-run the lens using the stored corpus
    from pramana.pipeline.orchestrator import _LENS_BY_NAME
    lens = _LENS_BY_NAME.get(request.lens_name)
    if lens is None:
        raise HTTPException(404, f"Unknown lens: {request.lens_name}")

    # We don't have corpus/evidence objects in memory after completion.
    # Rebuild a minimal NormalizedEvidence from DB facts.
    from pramana.pipeline.corpus import Corpus
    from pramana.pipeline.extraction import ExtractedFact
    from pramana.pipeline.normalization import NormalizedEvidence

    with get_session(settings) as session:
        db_papers = (
            session.query(Paper)
            .join(ExtractedFactDB, ExtractedFactDB.paper_id == Paper.id)
            .distinct().all()
        )
        db_facts = session.query(ExtractedFactDB).all()

        corpus_papers = []
        for p in db_papers:
            authors = json.loads(p.authors) if p.authors else []
            corpus_papers.append({
                "db_id": p.id, "title": p.title, "authors": authors,
                "year": p.year, "venue": p.venue or "", "abstract": p.abstract or "",
            })

        facts = [
            ExtractedFact(
                fact_type=f.fact_type, content=f.content,
                direct_quote=f.direct_quote, location=f.location,
                paper_id=f.paper_id, confidence=f.confidence or 0.0,
            )
            for f in db_facts
        ]

    corpus = Corpus(papers=corpus_papers)
    evidence = NormalizedEvidence(facts=facts)

    try:
        result = lens.analyze(corpus, evidence, query, settings)
    except Exception as e:
        raise HTTPException(500, f"Lens re-run failed: {e}")

    # Patch the stored result
    new_lr = {"lens": result.lens_name, "title": result.title,
               "summary": result.summary, "content": result.content}
    updated = False
    for i, lr in enumerate(store["result"]["lens_results"]):
        if lr["lens"] == request.lens_name:
            store["result"]["lens_results"][i] = new_lr
            updated = True
            break
    if not updated:
        store["result"]["lens_results"].append(new_lr)

    return new_lr


class SearchMoreRequest(BaseModel):
    query: str
    max_papers: int = 10


@app.post("/api/analyze/{run_id}/search-more")
async def search_more_papers(run_id: str, request: SearchMoreRequest):
    """Fetch additional papers for a follow-up query and add to the run's evidence."""
    if run_id not in _analysis_store:
        raise HTTPException(404, "Analysis run not found")
    store = _analysis_store[run_id]
    if store["status"] != "completed":
        raise HTTPException(400, "Run must be completed to search for more papers")

    settings = get_settings()
    from pramana.pipeline.corpus import build_corpus
    from pramana.pipeline.extraction import extract_all_evidence
    from pramana.pipeline.hypothesis import parse_hypothesis

    # Parse the follow-up query as a mini hypothesis
    mini_query = parse_hypothesis(request.query, "new", settings)
    mini_query.search_queries = [request.query]  # keep it simple, single direct query

    # Fetch small corpus
    new_corpus = build_corpus(mini_query, max_papers=request.max_papers, settings=settings)
    added = len([p for p in new_corpus.papers if not p.get("screened_out")])

    if added == 0:
        return {"added_papers": 0, "message": "No new papers found"}

    # Extract evidence
    new_facts = extract_all_evidence(new_corpus, mini_query, settings)

    # Update paper count in stored report
    report_data = store["result"]
    for lr in report_data.get("lens_results", []):
        if lr.get("lens") == "evidence_table":
            c = lr.get("content", {})
            c["papers_with_evidence"] = c.get("papers_with_evidence", 0) + added
            c["total_facts"] = c.get("total_facts", 0) + len(new_facts)
            break

    return {"added_papers": added, "new_facts": len(new_facts),
            "message": f"Added {added} papers with {len(new_facts)} new facts"}


# --- Batch J: Onboarding / Explore endpoints ---

class ExploreFieldRequest(BaseModel):
    field: str


class BuildHypothesisRequest(BaseModel):
    population: str
    intervention: str
    comparison: str = ""
    outcome: str
    domain: str = ""


class SuggestHypothesesRequest(BaseModel):
    field: str
    selected_titles: list[str]


@app.post("/api/explore/sample-papers")
async def explore_sample_papers(request: ExploreFieldRequest):
    """Fetch a handful of representative papers for a field to help users explore."""
    settings = get_settings()
    from pramana.pipeline.corpus import build_corpus
    from pramana.pipeline.hypothesis import parse_hypothesis
    mini_query = parse_hypothesis(request.field, "new", settings)
    mini_query.search_queries = [request.field]
    corpus = build_corpus(mini_query, max_papers=8, settings=settings)
    papers = [
        {"title": p.get("title", ""), "abstract": (p.get("abstract") or "")[:300],
         "year": p.get("year")}
        for p in corpus.papers[:8] if not p.get("screened_out")
    ]
    return {"papers": papers}


@app.post("/api/explore/suggest-hypotheses")
async def explore_suggest_hypotheses(request: SuggestHypothesesRequest):
    """Suggest research hypotheses based on a field and selected papers."""
    settings = get_settings()
    from pramana.llm.client import call_llm
    from pramana.llm.prompts import SUGGEST_HYPOTHESES
    titles_text = "\n".join(f"- {t}" for t in request.selected_titles)
    prompt = SUGGEST_HYPOTHESES.format(field=request.field, paper_titles=titles_text)
    raw = call_llm(
        system="You are a research methodology expert.",
        user=prompt,
        settings=settings,
    )
    import json as _json
    try:
        data = _json.loads(raw)
        return {"hypotheses": data.get("hypotheses", [])}
    except Exception:
        return {"hypotheses": [raw]}


@app.post("/api/explore/build-hypothesis")
async def explore_build_hypothesis(request: BuildHypothesisRequest):
    """Compose a hypothesis from PICO components."""
    settings = get_settings()
    import json as _json

    from pramana.llm.client import call_llm
    from pramana.llm.prompts import PICO_TO_HYPOTHESIS
    prompt = PICO_TO_HYPOTHESIS.format(
        population=request.population,
        intervention=request.intervention,
        comparison=request.comparison or "none specified",
        outcome=request.outcome,
        domain=request.domain or "not specified",
    )
    raw = call_llm(system="You are a research methodology expert.", user=prompt, settings=settings)
    try:
        data = _json.loads(raw)
        return {"hypothesis": data.get("hypothesis", raw)}
    except Exception:
        return {"hypothesis": raw}


def _bibtex_key(authors: list[str], year: int | None, title: str) -> str:
    """Generate a BibTeX citation key."""
    import re
    first_author = authors[0].split()[-1] if authors else "unknown"
    first_author = re.sub(r"[^a-zA-Z]", "", first_author).lower()
    year_str = str(year) if year else "xxxx"
    first_word = re.sub(r"[^a-zA-Z]", "", title.split()[0]).lower() if title else "paper"
    return f"{first_author}{year_str}{first_word}"


def _build_report_summary(report_data: dict, max_chars: int = 8000) -> str:
    """Build a text summary of a report for chat context."""
    lines: list[str] = []

    hyp = report_data.get("hypothesis", {})
    if hyp.get("topics"):
        lines.append(f"Topics: {', '.join(hyp['topics'])}")
    if hyp.get("domains"):
        lines.append(f"Domains: {', '.join(hyp['domains'])}")
    if hyp.get("methods"):
        lines.append(f"Methods: {', '.join(hyp['methods'])}")
    lines.append("")

    for lr in report_data.get("lens_results", []):
        lines.append(f"### {lr.get('title', lr.get('lens', ''))}")
        lines.append(lr.get("summary", ""))
        content = lr.get("content", {})

        if lr.get("lens") == "evidence_table":
            total = content.get("total_facts", 0)
            papers = content.get("papers_with_evidence", 0)
            lines.append(f"Total facts: {total}, Papers: {papers}")
        elif lr.get("lens") == "gap_discovery":
            for gap in content.get("gaps", [])[:5]:
                lines.append(f"- Gap ({gap.get('severity', '?')}): {gap.get('description', '')}")
        elif lr.get("lens") == "research_planning":
            for d in content.get("directions", [])[:5]:
                area = d.get("area", d.get("direction", "")) if isinstance(d, dict) else str(d)
                lines.append(f"- Direction: {area}")
        lines.append("")

    result = "\n".join(lines)
    return result[:max_chars]



# --- Background task ---

def _run_analysis(run_id: str) -> None:
    """Execute the full analysis pipeline as a background task."""
    settings = get_settings()
    store = _analysis_store[run_id]

    try:
        store["status"] = "running"
        logger.info("[%s] Pipeline started: %s", run_id[:8], store["hypothesis"][:80])

        # Stage 1: Parse hypothesis
        store["stage"] = "parsing"
        store["progress"] = {"step": 1, "total": 7, "description": "Parsing hypothesis"}
        logger.info("[%s] Stage 1/7: Parsing hypothesis", run_id[:8])
        from pramana.pipeline.hypothesis import parse_hypothesis
        parsed = parse_hypothesis(
            store["hypothesis"], store["initiation_type"], settings,
            prior_research=store.get("prior_research", ""),
            declared_domain=store.get("domain", ""),
        )
        parsed.hypothesis_text = store["hypothesis"]
        store["progress"]["parsed"] = parsed.model_dump()
        logger.info("[%s] Parsed → %d domains, %d topics, %d queries",
                    run_id[:8], len(parsed.domains), len(parsed.topics), len(parsed.search_queries))

        # Pause 1: Hypothesis confirmation (5 min timeout → auto-continue)
        store["parsed_query"] = parsed.model_dump()
        store["status"] = "awaiting_confirmation"
        store["stage"] = "awaiting_confirmation"
        logger.info("[%s] Paused for hypothesis confirmation", run_id[:8])
        store["_confirm_event"].wait(timeout=300)
        store["status"] = "running"
        # Apply any edits the user made
        edits = store.get("_confirm_edits")
        if edits:
            parsed.domains = edits.get("domains", parsed.domains)
            parsed.topics = edits.get("topics", parsed.topics)
            parsed.search_queries = edits.get("search_queries", parsed.search_queries)
            logger.info("[%s] Hypothesis updated by user", run_id[:8])

        # Stage 2: Build corpus
        store["stage"] = "retrieval"
        store["progress"] = {"step": 2, "total": 7, "description": "Retrieving papers"}
        logger.info("[%s] Stage 2/7: Retrieving papers (max=%d)", run_id[:8], store["max_papers"])
        from pramana.pipeline.corpus import build_corpus
        corpus = build_corpus(parsed, max_papers=store["max_papers"], settings=settings)
        store["progress"]["papers_found"] = len(corpus.papers)
        store["progress"]["sources"] = {
            "s2": corpus.total_from_s2,
            "arxiv": corpus.total_from_arxiv,
            "pubmed": corpus.total_from_pubmed,
            "crossref": corpus.total_from_crossref,
        }
        logger.info(
            "[%s] Corpus: %d papers (S2=%d, arXiv=%d, PubMed=%d, CrossRef=%d)",
            run_id[:8], len(corpus.papers), corpus.total_from_s2,
            corpus.total_from_arxiv, corpus.total_from_pubmed, corpus.total_from_crossref,
        )

        # Stage 3: Screen papers
        store["stage"] = "screening"
        store["progress"] = {"step": 3, "total": 7, "description": "Screening papers for relevance"}
        logger.info("[%s] Stage 3/7: Screening %d papers", run_id[:8], len(corpus.papers))
        from pramana.pipeline.screening import screen_corpus
        corpus = screen_corpus(corpus, parsed, settings)
        screened_count = sum(1 for p in corpus.papers if p.get("screened_out"))
        passed_count = len(corpus.papers) - screened_count
        store["progress"]["papers_screened_out"] = screened_count
        store["progress"]["papers_passed"] = passed_count
        logger.info(
            "[%s] Screening: %d passed, %d filtered",
            run_id[:8], passed_count, screened_count,
        )

        # Pause 2: Paper curation (10 min timeout → auto-continue)
        store["corpus_papers"] = [
            {
                "db_id": p.get("db_id"),
                "title": p.get("title", ""),
                "authors": p.get("authors", [])[:3],
                "year": p.get("year"),
                "venue": p.get("venue", ""),
                "source": p.get("source", "unknown"),
                "screened_out": bool(p.get("screened_out")),
                "screening_reason": p.get("screening_reason", ""),
                "relevance_score": round(p.get("relevance_score", 0.0), 3),
            }
            for p in corpus.papers
        ]
        store["status"] = "awaiting_curation"
        store["stage"] = "awaiting_curation"
        logger.info("[%s] Paused for paper curation", run_id[:8])
        store["_curate_event"].wait(timeout=600)
        store["status"] = "running"
        # Apply user exclusions
        excluded_ids = store.get("_excluded_paper_ids", set())
        if excluded_ids:
            for p in corpus.papers:
                if p.get("db_id") in excluded_ids:
                    p["screened_out"] = True
                    p["screening_reason"] = "Excluded by user"
            removed = sum(1 for p in corpus.papers if p.get("db_id") in excluded_ids)
            logger.info("[%s] User excluded %d papers", run_id[:8], removed)
        passed_count = sum(1 for p in corpus.papers if not p.get("screened_out"))

        # Stage 4: Extract evidence
        store["stage"] = "extraction"
        store["progress"] = {
            "step": 4, "total": 7,
            "description": f"Extracting evidence (0 / {passed_count} papers)",
            "papers_processed": 0,
            "papers_total": passed_count,
            "facts_extracted": 0,
            "current_paper": "",
        }
        logger.info("[%s] Stage 4/7: Extracting evidence from %d papers",
                    run_id[:8], passed_count)

        def _extraction_progress(done: int, total: int, title: str) -> None:
            short = title[:60] + "…" if len(title) > 60 else title
            store["progress"]["papers_processed"] = done
            store["progress"]["papers_total"] = total
            store["progress"]["current_paper"] = short
            store["progress"]["description"] = (
                f"Extracting evidence ({done} / {total}) — {short}"
            )

        from pramana.pipeline.extraction import extract_all_evidence
        evidence = extract_all_evidence(corpus, parsed, settings,
                                        progress_callback=_extraction_progress)
        store["progress"]["facts_extracted"] = len(evidence)
        store["progress"]["description"] = (
            f"Extracted {len(evidence)} facts from {passed_count} papers"
        )
        logger.info("[%s] Extracted %d facts", run_id[:8], len(evidence))

        # Stage 5: Normalize
        store["stage"] = "normalization"
        store["progress"] = {"step": 5, "total": 7, "description": "Normalizing evidence"}
        logger.info("[%s] Stage 5/7: Normalizing %d facts", run_id[:8], len(evidence))
        from pramana.pipeline.normalization import normalize_evidence
        normalized = normalize_evidence(evidence, settings)
        store["progress"]["mappings"] = len(normalized.canonical_mappings)
        store["progress"]["categories"] = len(normalized.categories)
        logger.info("[%s] Normalized: %d mappings, %d categories",
                    run_id[:8], len(normalized.canonical_mappings), len(normalized.categories))

        # Stage 6: Route to analysis flows and run lenses
        store["stage"] = "analysis"
        store["progress"] = {
            "step": 6, "total": 7, "description": "Routing and running analysis flows"
        }
        logger.info("[%s] Stage 6/7: Routing analysis flows", run_id[:8])
        from pramana.flows.router import select_flows
        from pramana.pipeline.orchestrator import run_flows
        action = store.get("action", "")
        selected_flows, routing_reasoning = select_flows(
            store["hypothesis"], action, parsed, settings
        )
        store["progress"]["selected_flows"] = [f.name for f in selected_flows]
        store["progress"]["routing_reasoning"] = routing_reasoning
        logger.info(
            "[%s] Flows selected: %s",
            run_id[:8], ", ".join(f.name for f in selected_flows),
        )
        results = run_flows(corpus, normalized, parsed, settings, selected_flows, routing_reasoning)
        store["progress"]["lenses_completed"] = results.active_lenses
        logger.info("[%s] Lenses completed: %s", run_id[:8], ", ".join(results.active_lenses))

        # Stage 7: Synthesize + generate report
        store["stage"] = "report"
        store["progress"] = {"step": 7, "total": 7, "description": "Synthesizing findings"}
        logger.info("[%s] Stage 7/7: Synthesizing and generating report", run_id[:8])
        from pramana.pipeline.orchestrator import synthesize_summary
        from pramana.report.generator import generate_report
        results.executive_summary = synthesize_summary(results, parsed, settings)
        store["progress"]["description"] = "Generating report"
        report_json = generate_report(results, parsed, "json", settings)

        store["status"] = "completed"
        store["stage"] = "done"
        store["result"] = json.loads(report_json)
        logger.info("[%s] Pipeline completed successfully", run_id[:8])

        # Update DB — persist the full report
        with get_session(settings) as session:
            db_run = session.get(AnalysisRun, store["db_run_id"])
            if db_run:
                db_run.status = "completed"
                db_run.completed_at = datetime.now(timezone.utc)
                db_run.results = report_json

    except Exception as e:
        store["status"] = "failed"
        store["error"] = str(e)
        logger.error("[%s] Pipeline FAILED at stage '%s': %s", run_id[:8], store["stage"], e)
        logger.debug("[%s] Traceback:\n%s", run_id[:8], traceback.format_exc())

        with get_session(settings) as session:
            db_run = session.get(AnalysisRun, store["db_run_id"])
            if db_run:
                db_run.status = "failed"
